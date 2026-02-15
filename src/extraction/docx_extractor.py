"""
Word document (.docx / .doc) text extraction.

Uses pandoc as primary extractor (preserves structure well),
with python-docx as fallback. For legacy .doc files, attempts
conversion via LibreOffice.
"""

import logging
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str | Path) -> tuple[str, str]:
    """
    Extract text from a Word document (.docx or .doc).

    Args:
        file_path: Path to the Word document.

    Returns:
        Tuple of (extracted_text, method_used).

    Raises:
        RuntimeError: If extraction fails with all methods.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    # Handle legacy .doc files by converting first
    if file_path.suffix.lower() == ".doc":
        return _handle_legacy_doc(file_path)

    # Try pandoc first
    text, method = _try_pandoc(file_path)
    if text:
        return text, method

    # Fallback to python-docx
    text, method = _try_python_docx(file_path)
    if text:
        return text, method

    raise RuntimeError(
        f"Could not extract text from {file_path}. "
        "Install pandoc or python-docx: pip install python-docx"
    )


def _try_pandoc(file_path: Path) -> tuple[str, str]:
    """Try extraction with pandoc (best structure preservation)."""
    try:
        result = subprocess.run(
            ["pandoc", str(file_path), "-t", "plain", "--wrap=none"],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            logger.info(
                f"Extracted {len(result.stdout)} chars from {file_path.name} via pandoc"
            )
            return result.stdout, "pandoc"
        elif result.stderr:
            logger.warning(f"pandoc warning for {file_path.name}: {result.stderr}")
    except FileNotFoundError:
        logger.debug("pandoc not available")
    except subprocess.TimeoutExpired:
        logger.warning(f"pandoc timed out on {file_path.name}")

    return "", ""


def _try_python_docx(file_path: Path) -> tuple[str, str]:
    """Try extraction with python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.debug("python-docx not available")
        return "", ""

    try:
        doc = Document(str(file_path))
        text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n\n".join(text_parts)
        if full_text.strip():
            logger.info(
                f"Extracted {len(full_text)} chars from {file_path.name} via python-docx"
            )
            return full_text, "python-docx"
        else:
            logger.warning(f"python-docx returned empty text for {file_path.name}")
            return "", ""
    except Exception as e:
        logger.warning(f"python-docx failed on {file_path.name}: {e}")
        return "", ""


def _handle_legacy_doc(file_path: Path) -> tuple[str, str]:
    """
    Handle legacy .doc files by converting to .docx via LibreOffice,
    then extracting text from the converted file.
    """
    logger.info(f"Legacy .doc file detected: {file_path.name}, attempting conversion")

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmpdir,
                    str(file_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode == 0:
                converted = Path(tmpdir) / (file_path.stem + ".docx")
                if converted.exists():
                    text, method = extract_text_from_docx(converted)
                    return text, f"libreoffice+{method}"

        logger.warning(f"LibreOffice conversion failed for {file_path.name}")
    except FileNotFoundError:
        logger.warning("LibreOffice (soffice) not available for .doc conversion")
    except subprocess.TimeoutExpired:
        logger.warning(f"LibreOffice conversion timed out for {file_path.name}")

    # Last resort: try pandoc directly on .doc
    text, method = _try_pandoc(file_path)
    if text:
        return text, method

    raise RuntimeError(
        f"Could not extract text from legacy .doc file: {file_path}. "
        "Install LibreOffice for .doc support."
    )
